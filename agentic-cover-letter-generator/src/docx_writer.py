"""Write the final cover letter draft to a .docx file."""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def _sanitize(text: str) -> str:
    """Lowercase text and strip it down to filename-safe characters.

    Whitespace becomes underscores; anything that isn't alphanumeric or an
    underscore afterward is dropped.
    """
    text = text.strip().lower()
    text = re.sub(r"\s+", "_", text)
    text = re.sub(r"[^a-z0-9_]", "", text)
    return text


def _add_hyperlink(paragraph, url: str, text: str, color: str = "000000") -> None:
    """Append a real, clickable hyperlink run to `paragraph`.

    python-docx has no public API for hyperlinks — this builds the
    relationship + <w:hyperlink> XML element directly, which is the
    standard workaround. `color` is a 6-digit hex string with no '#'
    (default black, so links don't render in Word's default blue).
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")

    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), color)
    run_props.append(color_elem)

    underline_elem = OxmlElement("w:u")
    underline_elem.set(qn("w:val"), "single")
    run_props.append(underline_elem)

    run.append(run_props)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


def _add_paragraph_with_markdown_bold(document: Document, text: str, alignment=None):
    """Add a paragraph, rendering **bold** markdown spans as real bold runs.

    The Generator's output uses **bold** markers for skill-paragraph
    headers (per instructions.md's own format spec); python-docx has no
    markdown awareness, so without this the literal asterisks would show
    up in the saved file instead of actual bold text.
    """
    para = document.add_paragraph()
    if alignment is not None:
        para.alignment = alignment

    pos = 0
    for match in _BOLD_PATTERN.finditer(text):
        if match.start() > pos:
            para.add_run(text[pos : match.start()])
        bold_run = para.add_run(match.group(1))
        bold_run.bold = True
        pos = match.end()
    if pos < len(text):
        para.add_run(text[pos:])

    return para


def _add_header(
    document: Document,
    full_name: str,
    location: str,
    phone: str,
    email: str,
    linkedin_url: str,
    github_url: str,
    portfolio_url: str,
) -> None:
    """Add the centered name + contact-line header block to `document`."""
    name_para = document.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(full_name)
    name_run.font.size = Pt(15)

    contact_para = document.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.add_run(f"{location} | {phone} | {email} | ")
    _add_hyperlink(contact_para, linkedin_url, "LinkedIn")
    contact_para.add_run(" | ")
    _add_hyperlink(contact_para, github_url, "GitHub")
    contact_para.add_run(" | ")
    _add_hyperlink(contact_para, portfolio_url, "Portfolio")


def write_docx(
    letter_text: str,
    company: str,
    role: str,
    output_dir: str = "output",
    *,
    full_name: str | None = None,
    location: str | None = None,
    phone: str | None = None,
    email: str | None = None,
    linkedin_url: str | None = None,
    github_url: str | None = None,
    portfolio_url: str | None = None,
) -> str:
    """Write `letter_text` to a .docx file in `output_dir`.

    Produces a business-letter document: 1-inch margins, single-spaced
    11pt Times New Roman body text, no header/footer section. `letter_text`
    is split into left-aligned paragraphs on blank lines; any `**bold**`
    markdown spans within a paragraph are rendered as real bold runs.

    If `full_name` is provided, a centered header is added above the body:
    the name at 15pt, then a centered contact line at 11pt with LinkedIn/
    GitHub/Portfolio rendered as black (non-default-blue) hyperlinks. Pass
    `full_name` and the rest of the contact fields together, or omit all of
    them to skip the header entirely (e.g. for tests).

    Args:
        letter_text: The final, reviewer-approved cover letter body text.
            Paragraphs should be separated by a blank line.
        company: Company name, used to build the output filename.
        role: Role/job title, used to build the output filename.
        output_dir: Directory to write the .docx file into. Created if it
            doesn't already exist.
        full_name: Full name for the header. Omit to skip the header block.
        location: City/state line for the contact line.
        phone: Phone number for the contact line.
        email: Email address for the contact line.
        linkedin_url: LinkedIn profile URL, rendered as a hyperlink.
        github_url: GitHub profile URL, rendered as a hyperlink.
        portfolio_url: Portfolio URL, rendered as a hyperlink.

    Returns:
        The full filepath of the saved .docx file, as a string.
    """
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    filename = f"{_sanitize(company)}_{_sanitize(role)}.docx"
    filepath = out_dir / filename

    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(12)

    if full_name:
        _add_header(
            document,
            full_name=full_name,
            location=location,
            phone=phone,
            email=email,
            linkedin_url=linkedin_url,
            github_url=github_url,
            portfolio_url=portfolio_url,
        )

    paragraphs = re.split(r"\n\s*\n", letter_text.strip())
    for para_text in paragraphs:
        para_text = para_text.strip()
        if para_text:
            _add_paragraph_with_markdown_bold(
                document, para_text, alignment=WD_ALIGN_PARAGRAPH.LEFT
            )

    document.save(str(filepath))
    return str(filepath)


if __name__ == "__main__":
    sample_letter = """Dear Recruiter,

My name is Jane Doe, and I am writing to express my interest in the Software Engineer position at Test Company. I am currently studying Computer Science, with a strong focus on distributed systems and developer tooling. I believe the following skills make me a strong fit for this role.

**Backend Systems:** During my internship at a fintech startup, I designed and shipped a rate-limiting service that reduced downstream API errors by 40%, working primarily in Python and Go.

**Distributed Systems Coursework:** My graduate coursework in distributed systems covered consensus protocols, replication, and fault tolerance, which I applied in a course project building a Raft-based key-value store.

**Open Source Contributions:** I have contributed several merged pull requests to a widely-used open-source observability tool, focused on improving its tracing instrumentation.

As for my interest in Test Company, it is deeply rooted in the company's engineering culture and its investment in developer tooling. I've followed the team's engineering blog closely and admire the emphasis on internal platform quality. I am enthusiastic about the possibility of bringing my skills to Test Company and look forward to discussing how my background and experience would allow me to contribute to the Platform team.

Sincerely,

Jane Doe"""

    saved_path = write_docx(
        letter_text=sample_letter,
        company="Test Company",
        role="Software Engineer",
        full_name="Jane Doe",
        location="San Francisco, CA",
        phone="+1 (555) 555-5555",
        email="jane.doe@example.com",
        linkedin_url="https://www.linkedin.com/in/janedoe",
        github_url="https://github.com/janedoe",
        portfolio_url="https://janedoe.dev",
    )
    print(f"Saved test letter to: {saved_path}")
